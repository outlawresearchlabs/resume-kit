"""
docx_builder.py — generate an ATS-safe DOCX resume.

ATS systems prefer DOCX because the XML structure gives a predictable
extraction order. This builder intentionally uses:
- Single-column layout
- Standard section headers
- Month YYYY – Present date format
- Bare, comma-separated skills tokens
- No tables, text boxes, multi-column, or decorative Unicode
"""
import os
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


def _norm_date(dates: str) -> str:
    """Convert MM/YYYY or 06/2014 - Current to Month YYYY – Present."""
    if not dates:
        return ""
    # Replace common variants
    text = dates.replace("–", "-").replace("Current", "Present").replace("current", "Present")
    # Look for MM/YYYY or M/YYYY patterns
    def repl(m):
        mm = int(m.group(1))
        yyyy = m.group(2)
        if 1 <= mm <= 12:
            return f"{MONTHS[mm - 1]} {yyyy}"
        return m.group(0)
    text = re.sub(r"\b(0?[1-9]|1[0-2])/(\d{4})\b", repl, text)
    return text


def _clean_paragraph(p):
    """Remove space-after from a paragraph for tight lists."""
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    _clean_paragraph(p)
    return p


def _add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    _clean_paragraph(p)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    _clean_paragraph(p)
    return p


def _contact_block(doc, contact):
    name = contact.get("name", "").strip()
    if name:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clean_paragraph(p)

    pieces = []
    for key in ("phone", "email", "location", "linkedin", "github"):
        val = contact.get(key, "").strip()
        if val:
            pieces.append(val)
    if pieces:
        p = doc.add_paragraph(" · ".join(pieces))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clean_paragraph(p)

    doc.add_paragraph()  # small gap


def _job_bullets(job):
    """Return non-empty bullet strings in dict order."""
    bullets = []
    for k, v in job.items():
        if k in ("role", "company", "dates") or not isinstance(v, str):
            continue
        if v.strip():
            bullets.append(v.strip())
    return bullets


def build_ats_docx(profile_data, config_key="executive", overrides=None, out_path: Optional[str] = None):
    """Build an ATS-safe DOCX from profile data and return the output path."""
    overrides = overrides or {}

    contact = profile_data.get("contact", {})
    jobs_raw = profile_data.get("jobs", {})
    if isinstance(jobs_raw, dict):
        jobs = []
        for key in ("current", "previous", "early"):
            if jobs_raw.get(key):
                jobs.append(jobs_raw[key])
    else:
        jobs = jobs_raw

    skills = profile_data.get("skills", {})
    sections = profile_data.get("sections", {})

    summary = overrides.get("summary") or sections.get("summary", "")

    # Skill order from override or config heuristic.
    skill_order = overrides.get("skill_order")
    if not skill_order:
        if config_key == "executive":
            skill_order = ["leadership", "security", "risk_compliance", "cloud_infra", "scripting", "customer_facing"]
        elif config_key == "technical":
            skill_order = ["security", "scripting", "cloud_infra", "networking", "leadership", "customer_facing"]
        else:
            skill_order = ["customer_facing", "security", "cloud_infra", "networking", "scripting", "leadership"]

    doc = Document()
    sections_doc = doc.sections[0]
    sections_doc.top_margin = Inches(0.7)
    sections_doc.bottom_margin = Inches(0.7)
    sections_doc.left_margin = Inches(0.8)
    sections_doc.right_margin = Inches(0.8)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    _contact_block(doc, contact)

    if summary:
        _add_heading(doc, "Summary")
        _add_paragraph(doc, summary)
        doc.add_paragraph()

    # Flat skills section for ATS.
    all_skills = []
    for cat in skill_order:
        if cat in skills and skills[cat]:
            all_skills.extend([s.strip() for s in skills[cat].split(",") if s.strip()])
    if all_skills:
        _add_heading(doc, "Skills")
        _add_paragraph(doc, ", ".join(all_skills))
        doc.add_paragraph()

    if jobs:
        _add_heading(doc, "Work Experience")
        for job in jobs:
            role = job.get("role", "").strip()
            company = job.get("company", "").strip()
            dates = _norm_date(job.get("dates", ""))

            header = role
            if company:
                header = f"{role}, {company}"
            if dates:
                header = f"{role}, {company} — {dates}"
            p = doc.add_paragraph()
            r = p.add_run(header)
            r.bold = True
            _clean_paragraph(p)

            bullets = _job_bullets(job)
            for b in bullets:
                _add_bullet(doc, b)
            doc.add_paragraph()

    education = sections.get("education", [])
    if education:
        _add_heading(doc, "Education")
        for degree, school, dates in education:
            line = f"{degree}, {school}"
            if dates:
                line = f"{degree}, {school} — {_norm_date(dates)}"
            _add_paragraph(doc, line)
        doc.add_paragraph()

    certs = sections.get("certs", [])
    if certs:
        _add_heading(doc, "Certifications")
        for c in certs:
            _add_bullet(doc, c)
        doc.add_paragraph()

    awards = sections.get("awards", [])
    if awards:
        _add_heading(doc, "Awards")
        for a in awards:
            _add_bullet(doc, a)
        doc.add_paragraph()

    publications = []
    pub_raw = sections.get("publication")
    if isinstance(pub_raw, dict):
        if any(str(v).strip() for v in pub_raw.values() if isinstance(v, str)):
            publications.append(pub_raw)
    elif isinstance(pub_raw, list):
        publications = [p for p in pub_raw if any(str(v).strip() for v in (p.values() if isinstance(p, dict) else p) if isinstance(v, str))]
    if publications:
        _add_heading(doc, "Publications")
        for pub in publications:
            title = pub.get("title", "") if isinstance(pub, dict) else pub[0]
            _add_paragraph(doc, title)
        doc.add_paragraph()

    if not out_path:
        safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", contact.get("name", "resume")).strip("_").lower() or "resume"
        out_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            "out",
            f"{safe_name}_{config_key}_ats.docx",
        )
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path
